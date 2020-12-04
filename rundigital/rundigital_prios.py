# rundigital_prios.py - web scrapping

"""Lister uniquement les sujets, dans les 18 boîtes qu'on a retenu."""

import json
from lxml import etree, html
from docx_common import add_char_run, add_pinyin_run, add_text_run, add_html_run, add_paragraph, add_heading

# Writing out to Word .docx files
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_LINE_SPACING

#-------------------------------------------------------------------------------
# dump_text
#-------------------------------------------------------------------------------

prios = [
    'STORENGY',
    'ALGECO',
    'DELPHI TECHNOLOGIES',
    'TOTAL DIRECT ENERGIE',
    'ACIAL',
    'APRC',
    'GALIAN',
    'FCA FRANCE',
    'ILE DE FRANCE MOBILITE',
    'INNOTHERA',
    'POUEY INTERNATIONAL',
    'RTE',
    'STPI SOCIETE TECHNIQUE DE PRODUCTIONS INDUSTRIELLES',
    'CARREFOUR FRANCE',
    'SMARTADSERVER',
    'ALKAN',
    'GRT GAZ',
    'MAN TRUCK & BUS FRANCE',
    'HERMES PARFUMS',
    'TELEDYNE E2V',
]

#-------------------------------------------------------------------------------
# get_prio_sujets
#-------------------------------------------------------------------------------

def get_prio_sujets(sociétés):
    d = {}
    for société in sociétés:
        if société['nom'] not in prios:
            continue

        # Sections
        sections = []
        for sect in société['sections']:
            dd = {}
            if not sect['section'].startswith('Quel sujet'):
                continue
            dd['sujet'] = sect['section']
            dd['questions'] = sect['questions']
            sections.append(dd)

            # foundq = False
            # for q in sect['questions']:
            #     if q['question'].startswith('Détails de vos sujets'):
            #         dd['texte'] = q['réponse']
            #         foundq = True
            #         break

            # if foundq:
            #     sections.append(dd)

        if len(sections) > 0:
            d[société['nom']] = sections

    return d

#-------------------------------------------------------------------------------
# get_sociétés
#-------------------------------------------------------------------------------

def get_sociétés(fp):
    with open(fp, 'r', encoding='utf-8') as f:
        obj = f.read()

#-------------------------------------------------------------------------------
# write_docx
#-------------------------------------------------------------------------------

def write_docx(obj):
    # Output a .docx file in the current directory
    outpath = r'C:\a\rundigital\priorités.docx'

    # Write out the contents as Word
    doc = Document('empty.docx')

    # Set normal margins
    s = doc.sections[0]
    s.left_margin = Inches(0.59)
    s.right_margin = Inches(0.59)
    s.top_margin = Inches(0.59)
    s.bottom_margin = Inches(0.59)

    for k, v in obj.items():
        # Nom de la boîte
        add_heading(doc, 1, k)

        for sect in v:
            # Numéro du sujet
            add_heading(doc, 2, sect['sujet'])

            for q in sect['questions']:
                add_heading(doc, 3, q['question'])
                
                # One paragraph for everyting
                p = doc.add_paragraph()

                pf = p.paragraph_format
                pf.space_before = Pt(4)
                pf.space_after = Pt(4)
                pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

                add_text_run(p, q['réponse'])

    doc.save(outpath)

#-------------------------------------------------------------------------------
# main
#-------------------------------------------------------------------------------

filename = r'C:\a\rundigital\rundigital.json'
with open(filename, 'r', encoding='utf-8') as f:
    obj = json.load(f)

result = get_prio_sujets(obj)

# Write out json file
filename = r'C:\a\rundigital\priorités.json'
print(f'Writing {filename}')
with open(filename, 'w', encoding='utf-8') as f:
    json.dump(result, f, indent=4, ensure_ascii=False)

# Write out Word .docx file
write_docx(result)