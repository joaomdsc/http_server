# office.py - produce word and excel versions of the RunDigital data

import re
import json
from lxml import etree, html
from docx_common import add_char_run, add_pinyin_run, add_text_run, add_html_run, add_paragraph, add_heading

# Writing out to Word .docx files
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_LINE_SPACING

# Writing out to Excel .xlsx files
from openpyxl import Workbook
from openpyxl.styles import Alignment
from openpyxl.utils import get_column_letter

#-------------------------------------------------------------------------------
# scontact
#-------------------------------------------------------------------------------

def scontact(contact):
    x = []
    for c in contact:
        x.extend(c)
    return '\n'.join(x)

#-------------------------------------------------------------------------------
# write_docx
#-------------------------------------------------------------------------------

def write_docx(obj):
    # Output a .docx file in the current directory
    outpath = 'rundigital.docx'

    # Write out the contents as Word
    doc = Document('empty.docx')

    # Set normal margins
    s = doc.sections[0]
    s.left_margin = Inches(0.59)
    s.right_margin = Inches(0.59)
    s.top_margin = Inches(0.59)
    s.bottom_margin = Inches(0.59)

    for o in obj:
        # Nom de la boîte
        add_heading(doc, 1, o['nom'])

        # Infos société
        p = doc.add_paragraph()
        add_text_run(p, f'CA: ', bold=True)
        add_text_run(p, f'{o["CA"]}\n')
        add_text_run(p, f'Effectif: ', bold=True)
        add_text_run(p, f'{o["Effectif"]}\n')

        # Some companies (e.g. FCA FRANCE) were not in the second web site
        s = f'{o["ape_code"]} ' if 'ape_code' in o else ''
        add_text_run(p, f'Code APE: ', bold=True)
        add_text_run(p, f'{s}{o["Code ape"]}\n')
        
        add_text_run(p, f'Activité: ', bold=True)
        add_text_run(p, f'{o["Activité"]}\n')

        if 'adresse' in o:
            add_text_run(p, f'Adresse:\n', bold=True)
            for s in  o['adresse']:
                add_text_run(p, f'    {s}\n')
        if 'tels' in o:
            stels = '\n'.join(o['tels'])
            add_text_run(p, f'Tels:\n', bold=True)
            for s in o['tels']:
                add_text_run(p, f'    {s}\n')
            
        add_text_run(p, f'Région: ', bold=True)
        add_text_run(p, f'{o["Région"]}\n')
        if 'url' in o:
            add_text_run(p , f'URL: ', bold=True)
            add_text_run(p , f'{o["url"]}\n')
        add_text_run(p, f'Fonctions des participants: ', bold=True)
        add_text_run(p, f'{o["Fonctions des participants"]}\n')
        if 'contact' in o:
            add_text_run(p, f'Contact:\n', bold=True)
            add_text_run(p, f'{scontact(o["contact"])}')

        for sect in o['sections']:
            # Numéro du sujet
            add_heading(doc, 2, sect['section'])

            for q in sect['questions']:
                add_heading(doc, 3, q['question'])
                
                # One paragraph for everyting
                p = doc.add_paragraph()

                pf = p.paragraph_format
                pf.space_before = Pt(4)
                pf.space_after = Pt(4)
                pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

                add_text_run(p, q['réponse'])

    doc.save(outpath)

#-------------------------------------------------------------------------------
# write_xlsx
#-------------------------------------------------------------------------------

def write_xlsx(obj):
    # Output a .xlsx file in the current directory
    outpath = 'rundigital.xlsx'

    wb = Workbook()
    ws = wb.active
    ws.title = 'RunDigital 2020'

    # Headers
    headers = [
        'Société',
        'Adresse',
        'Région',
        'Tels',
        'Fonctions Participants',
        'Contact',
        'Url',
        'CA',
        'Effectif',
        'Code APE',
        'Texte APE',
        'Activité',
        'Cadre1',
        'Sujet1',
        'C.Charges1',  # parfois Contraintes
        'Contexte1',
        'Cadre2',
        'Sujet2',
        'C.Charges2',
        'Contexte2',
        'Cadre3',
        'Sujet3',
        'C.Charges3',
        'Contexte3',
    ]
    for i, hdr in enumerate(headers):
        ws.cell(row=1, column=i+1, value=hdr)

    # One row per company
    row = 2
    for o in obj:
        # Nom de la boîte
        ws.cell(row=row, column=1, value=o['nom'])
        print(o['nom'])

        # Infos société
        if 'adresse' in o:
            c = ws.cell(row=row, column=2, value='\n'.join(o['adresse']))
            c.alignment = Alignment(wrapText=True)
        ws.cell(row=row, column=3, value=o['Région'])
        if 'tels' in o:
            c = ws.cell(row=row, column=4, value='\n'.join(o['tels']))
            c.alignment = Alignment(wrapText=True)
        ws.cell(row=row, column=5, value=o['Fonctions des participants'])
        if 'contact' in o:
            c = ws.cell(row=row, column=6, value=scontact(o["contact"]))
            c.alignment = Alignment(wrapText=True)
        if 'url' in o:
            ws.cell(row=row, column=7, value=o['url'])
        ws.cell(row=row, column=8, value=o['CA'])
        ws.cell(row=row, column=9, value=o['Effectif'])
        if 'ape_code' in o:
            ws.cell(row=row, column=10, value=o['ape_code'])
        ws.cell(row=row, column=11, value=o['Code ape'])
        ws.cell(row=row, column=12, value=o['Activité'])
            
        for sect in o['sections']:
            m = re.match('Quel sujet n°([1-3])', sect['section'])
            if not m:
                continue
            i = int(m.group(1))
            questions = sect['questions']
            ws.cell(row=row, column=4*(i-1) + 13, value=questions[0]['réponse'])
            if len(questions) >= 2:
                ws.cell(row=row, column=4*(i-1) + 14, value=questions[1]['réponse'])
            if len(questions) >= 3:
                ws.cell(row=row, column=4*(i-1) + 15, value=questions[2]['réponse'])
            if len(questions) >= 4:
                ws.cell(row=row, column=4*(i-1) + 16, value=questions[3]['réponse'])
        row += 1

    wb.save(filename=outpath)

#-------------------------------------------------------------------------------
# main
#-------------------------------------------------------------------------------

filename = 'rundigital_merged.json'
with open(filename, 'r', encoding='utf-8') as f:
    obj = json.load(f)

# Write out Word .docx file
# write_docx(obj)
write_xlsx(obj)